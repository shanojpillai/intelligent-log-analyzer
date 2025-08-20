import os
import redis
import json
import logging
import time
from datetime import datetime
from pathlib import Path
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from sqlalchemy import create_engine
from typing import Dict, Any

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.redis_client = redis.Redis.from_url(os.getenv("REDIS_URL", "redis://:redis_password_123@redis:6379"))
        self.db_engine = create_engine(os.getenv("DATABASE_URL", "postgresql://postgres:secure_password_123@postgres:5432/log_analyzer"))
        self.export_dir = Path("/app/exports")
        self.export_dir.mkdir(exist_ok=True)
        
    def compile_final_results(self, job_id: str) -> Dict[str, Any]:
        """Compile all analysis results into final report data"""
        try:
            extraction_data = self.redis_client.get(f"extraction:{job_id}")
            embedding_data = self.redis_client.get(f"embeddings:{job_id}")
            retrieval_data = self.redis_client.get(f"retrieval:{job_id}")
            ai_analysis_data = self.redis_client.get(f"ai_analysis:{job_id}")
            nlu_data = self.redis_client.get(f"nlu:{job_id}")
            
            extraction = json.loads(extraction_data.decode()) if extraction_data else {}
            embeddings = json.loads(embedding_data.decode()) if embedding_data else {}
            retrieval = json.loads(retrieval_data.decode()) if retrieval_data else {}
            ai_analysis = json.loads(ai_analysis_data.decode()) if ai_analysis_data else {}
            nlu = json.loads(nlu_data.decode()) if nlu_data else {}
            
            final_results = {
                "job_id": job_id,
                "generated_at": datetime.utcnow().isoformat(),
                "summary": {
                    "files_processed": extraction.get("total_files", 0),
                    "issues_found": len(ai_analysis.get("recommendations", [])),
                    "confidence": ai_analysis.get("confidence", 0),
                    "severity": ai_analysis.get("severity", "UNKNOWN")
                },
                "extraction_results": extraction,
                "embedding_results": embeddings,
                "similar_cases": retrieval.get("similar_cases", []),
                "historical_solutions": retrieval.get("historical_solutions", []),
                "ai_analysis": ai_analysis,
                "nlu_analysis": nlu,
                "key_findings": ai_analysis.get("key_findings", []),
                "recommendations": ai_analysis.get("recommendations", []),
                "entities": nlu.get("entities", []),
                "keywords": nlu.get("keywords", []),
                "log_patterns": nlu.get("log_patterns", [])
            }
            
            return final_results
            
        except Exception as e:
            logger.error(f"Error compiling results: {e}")
            return {"error": str(e)}
            
    def generate_pdf_report(self, job_id: str, results: Dict[str, Any]) -> str:
        """Generate PDF report"""
        try:
            filename = f"log_analysis_report_{job_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.export_dir / filename
            
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph("Intelligent Log Analysis Report", title_style))
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary = results.get("summary", {})
            summary_text = f"""
            <b>Job ID:</b> {job_id}<br/>
            <b>Generated:</b> {results.get('generated_at', 'N/A')}<br/>
            <b>Files Processed:</b> {summary.get('files_processed', 0)}<br/>
            <b>Issues Found:</b> {summary.get('issues_found', 0)}<br/>
            <b>Confidence Score:</b> {summary.get('confidence', 0)}%<br/>
            <b>Severity Level:</b> {summary.get('severity', 'UNKNOWN')}
            """
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Key Findings", styles['Heading2']))
            findings = results.get("key_findings", [])
            for i, finding in enumerate(findings, 1):
                story.append(Paragraph(f"{i}. {finding}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Recommendations", styles['Heading2']))
            recommendations = results.get("recommendations", [])
            for i, rec in enumerate(recommendations, 1):
                story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Similar Historical Cases", styles['Heading2']))
            similar_cases = results.get("similar_cases", [])[:3]  # Top 3
            if similar_cases:
                for i, case in enumerate(similar_cases, 1):
                    case_text = f"""
                    <b>Case {i}:</b><br/>
                    <b>Similarity:</b> {case.get('similarity', 0):.1f}%<br/>
                    <b>Error Type:</b> {case.get('error_type', 'N/A')}<br/>
                    <b>Content:</b> {case.get('content', 'N/A')[:200]}...
                    """
                    story.append(Paragraph(case_text, styles['Normal']))
                    story.append(Spacer(1, 10))
            else:
                story.append(Paragraph("No similar cases found.", styles['Normal']))
            
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Extracted Entities", styles['Heading2']))
            entities = results.get("entities", [])[:10]  # Top 10
            if entities:
                entity_data = [["Entity", "Type", "Confidence"]]
                for entity in entities:
                    entity_data.append([
                        entity.get("text", ""),
                        entity.get("label", ""),
                        f"{entity.get('confidence', 0):.2f}"
                    ])
                
                entity_table = Table(entity_data)
                entity_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(entity_table)
            else:
                story.append(Paragraph("No entities extracted.", styles['Normal']))
            
            doc.build(story)
            logger.info(f"Generated PDF report: {filename}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            return ""
            
    def generate_word_report(self, job_id: str, results: Dict[str, Any]) -> str:
        """Generate Word document report"""
        try:
            filename = f"log_analysis_report_{job_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.export_dir / filename
            
            doc = Document()
            
            title = doc.add_heading('Intelligent Log Analysis Report', 0)
            title.alignment = 1  # Center alignment
            
            doc.add_heading('Executive Summary', level=1)
            summary = results.get("summary", {})
            summary_table = doc.add_table(rows=6, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ("Job ID", job_id),
                ("Generated", results.get('generated_at', 'N/A')),
                ("Files Processed", str(summary.get('files_processed', 0))),
                ("Issues Found", str(summary.get('issues_found', 0))),
                ("Confidence Score", f"{summary.get('confidence', 0)}%"),
                ("Severity Level", summary.get('severity', 'UNKNOWN'))
            ]
            
            for i, (key, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = key
                summary_table.cell(i, 1).text = value
            
            doc.add_heading('Key Findings', level=1)
            findings = results.get("key_findings", [])
            for i, finding in enumerate(findings, 1):
                doc.add_paragraph(f"{i}. {finding}", style='List Number')
            
            doc.add_heading('Recommendations', level=1)
            recommendations = results.get("recommendations", [])
            for i, rec in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
            
            doc.add_heading('Similar Historical Cases', level=1)
            similar_cases = results.get("similar_cases", [])[:3]
            if similar_cases:
                for i, case in enumerate(similar_cases, 1):
                    doc.add_heading(f'Case {i}', level=2)
                    doc.add_paragraph(f"Similarity: {case.get('similarity', 0):.1f}%")
                    doc.add_paragraph(f"Error Type: {case.get('error_type', 'N/A')}")
                    doc.add_paragraph(f"Content: {case.get('content', 'N/A')[:200]}...")
            else:
                doc.add_paragraph("No similar cases found.")
            
            doc.add_heading('Extracted Entities', level=1)
            entities = results.get("entities", [])[:10]
            if entities:
                entity_table = doc.add_table(rows=1, cols=3)
                entity_table.style = 'Table Grid'
                
                hdr_cells = entity_table.rows[0].cells
                hdr_cells[0].text = 'Entity'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Confidence'
                
                for entity in entities:
                    row_cells = entity_table.add_row().cells
                    row_cells[0].text = entity.get("text", "")
                    row_cells[1].text = entity.get("label", "")
                    row_cells[2].text = f"{entity.get('confidence', 0):.2f}"
            else:
                doc.add_paragraph("No entities extracted.")
            
            doc.save(str(filepath))
            logger.info(f"Generated Word report: {filename}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            return ""
            
    def process_jobs(self):
        """Process export jobs from Redis queue"""
        logger.info("Starting export service")
        
        while True:
            try:
                job_data = self.redis_client.brpop("compilation_queue", timeout=5)
                
                if job_data:
                    job_id = job_data[1].decode()
                    logger.info(f"Processing export job: {job_id}")
                    
                    self.redis_client.hset(f"job:{job_id}", "status", "compiling_results")
                    
                    final_results = self.compile_final_results(job_id)
                    
                    if "error" not in final_results:
                        self.redis_client.set(f"results:{job_id}", json.dumps(final_results))
                        
                        pdf_path = self.generate_pdf_report(job_id, final_results)
                        
                        self.redis_client.hset(f"job:{job_id}", "status", "completed")
                        self.redis_client.hset(f"job:{job_id}", "progress", "100")
                        
                        logger.info(f"Completed processing for job {job_id}")
                    else:
                        self.redis_client.hset(f"job:{job_id}", "status", "failed")
                        self.redis_client.hset(f"job:{job_id}", "error", final_results.get("error", "Compilation failed"))
                        
            except Exception as e:
                logger.error(f"Error processing jobs: {e}")
                time.sleep(1)

if __name__ == "__main__":
    service = ExportService()
    service.process_jobs()
